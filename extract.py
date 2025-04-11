import fitz  # PyMuPDF
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re
import os

# Extract text from PDF
def extract_text_from_pdf(pdf_path):
    text = ""
    with fitz.open(pdf_path) as pdf:
        for page_num in range(pdf.page_count):
            page = pdf[page_num]
            page_text = page.get_text("text")
            
            # Filter out non-XML-compatible characters
            page_text = re.sub(r'[^\x20-\x7E]+', '', page_text)  # Remove non-ASCII characters
            
            text += page_text
            text += "\n\n"  # Separate pages with double newline
    return text

# Extract images from PDF
def extract_images_from_pdf(pdf_path, image_folder):
    with fitz.open(pdf_path) as pdf:
        image_list = []
        for page_num in range(pdf.page_count):
            page = pdf[page_num]
            images = page.get_images(full=True)
            for img_index, img in enumerate(images):
                xref = img[0]
                pix = fitz.Pixmap(pdf, xref)
                img_filename = f"{image_folder}/image_{page_num + 1}_{img_index + 1}.png"
                pix.save(img_filename)
                image_list.append(img_filename)
                pix = None  # Free memory
    return image_list

# Save extracted text into DOCX file
def save_text_in_docx(text, docx_path, image_folder):
    doc = Document()

    # APA style header
    header = doc.sections[0].header
    header_para = header.paragraphs[0]
    header_para.text = "Running head: TITLE OF YOUR DOCUMENT"  # Replace with your document title
    header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    # Title page formatting
    title = doc.add_paragraph("Title of Your Document")  # Replace with actual title
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].font.size = Pt(14)
    title.runs[0].bold = True
    title.runs[0].font.name = 'Times New Roman'

    # Author information
    author = doc.add_paragraph("Author Name\nInstitution Name")  # Replace with actual details
    author.alignment = WD_ALIGN_PARAGRAPH.CENTER
    author.runs[0].font.size = Pt(12)
    author.runs[0].font.name = 'Times New Roman'

    # Add a space before the main content
    doc.add_paragraph("\n")

    # Add extracted text with APA-style formatting
    paragraphs = text.split("\n\n")  # Separate text into paragraphs
    for para_text in paragraphs:
        para = doc.add_paragraph(para_text)
        para_format = para.paragraph_format
        para_format.line_spacing = 1.5  # Line spacing
        para_format.first_line_indent = Pt(18)  # Indentation
        if para.runs:
            para.runs[0].font.size = Pt(12)  # Body text size
            para.runs[0].font.name = 'Times New Roman'  # Set font for body text

    # Add references (placeholder, modify as necessary)
    doc.add_paragraph("\nReferences", style='Heading 1')
    doc.add_paragraph("Author Last Name, First Initial. (Year). Title of the Book. Publisher.")

    # Extract and insert images if any
    images = extract_images_from_pdf(pdf_path, image_folder)
    for image_path in images:
        doc.add_paragraph("Figure: ", style='Heading 2')
        doc.add_paragraph(f"See the image at {image_path}")
        doc.add_picture(image_path)

    # Save the document
    doc.save(docx_path)
    print(f"Document saved at {docx_path}")

# Main Execution
pdf_path = r'C:\Users\People\Desktop\upwork\PPT deep learning\heart_sound.pdf'  # Replace with your PDF path
docx_path = r'C:\Users\People\Desktop\upwork\PPT deep learning\extract_text.docx'  # Desired output path for the DOCX file
image_folder = r'C:\Users\People\Desktop\upwork\PPT deep learning\images'  # Folder to store extracted images

# Ensure image folder exists
os.makedirs(image_folder, exist_ok=True)

# Step 1: Extract text
extract_text = extract_text_from_pdf(pdf_path)

# Step 2: Save extracted text and images in APA-formatted DOCX
save_text_in_docx(extract_text, docx_path, image_folder)
