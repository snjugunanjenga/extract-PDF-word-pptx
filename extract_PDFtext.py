import fitz  # PyMuPDF
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re
import os

# Extract text from PDF
def extract_text_from_pdf(pdf_path):
    text = ""
    with fitz.open(pdf_path) as pdf:
        for page_num in range(pdf.page_count):
            page = pdf[page_num]
            page_text = page.get_text("text")
            
            # Filter out non-ASCII characters
            page_text = re.sub(r'[^\x20-\x7E]+', '', page_text)  # Clean text
            
            text += page_text
            text += "\n\n"  # Separate pages with double newline
    return text

# Save extracted text into a DOCX file
def save_text_in_docx(text, docx_path):
    doc = Document()

    # Add APA-style title
    doc.add_heading('Extracted Text', level=1).alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add text with APA formatting
    paragraphs = text.split("\n\n")  # Split text into paragraphs
    for para_text in paragraphs:
        if para_text.strip():  # Avoid adding empty paragraphs
            para = doc.add_paragraph(para_text.strip())
            para_format = para.paragraph_format
            para_format.line_spacing = 1.5  # Line spacing
            para_format.first_line_indent = Pt(18)  # Indentation
            if para.runs:
                para.runs[0].font.size = Pt(12)  # Font size
                para.runs[0].font.name = 'Times New Roman'  # Font type

    # Save document
    doc.save(docx_path)
    print(f"Document saved at {docx_path}")

# Main Execution
if __name__ == "__main__":
    # Replace with your PDF and DOCX paths
    pdf_path = r'C:\Users\People\Documents\MidlevelBackendDeveloperJD.pdf'
    docx_path = r'C:\Users\People\Documents\MidlevelBackendDeveloper.docx'

    # Step 1: Extract text from PDF
    extracted_text = extract_text_from_pdf(pdf_path)

    # Step 2: Save the extracted text in APA-formatted DOCX
    save_text_in_docx(extracted_text, docx_path)
